#!/usr/bin/env python3
"""Build the canonical CQI Word document template."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "brands/cqi/plantillas/CQI_DocumentTemplate_20260720.docx"
LOGO = ROOT / "brands/cqi/brand_style/CQISense_Design_System/assets/cqi-logo.png"
CONTENT_DXA = 9744
FONT_BODY = "Arial"
FONT_HEADING = "Arial"
FONT_NUMERIC = "Arial"


COLORS = {
    "brand": "1A43F5",
    "brand_hover": "0037D5",
    "navy": "001450",
    "accent": "FF8093",
    "repair": "2D7173",
    "risk": "FF563F",
    "operational": "37A781",
    "gray_50": "F7F7F7",
    "gray_100": "F0F0F0",
    "gray_200": "E1E1E1",
    "gray_600": "686868",
    "gray_800": "232324",
    "white": "FFFFFF",
}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name: str, size: int | float | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, color: str = "E1E1E1", size: str = "8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int, col_widths_dxa: list[int] | None = None):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.autofit = False
    if col_widths_dxa:
        grid = table._tbl.tblGrid
        if grid is None:
            grid = OxmlElement("w:tblGrid")
            table._tbl.insert(0, grid)
        for child in list(grid):
            grid.remove(child)
        for width in col_widths_dxa:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for idx, width in enumerate(col_widths_dxa):
                row.cells[idx].width = width
                tc_w = row.cells[idx]._tc.get_or_add_tcPr().find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    row.cells[idx]._tc.get_or_add_tcPr().append(tc_w)
                tc_w.set(qn("w:w"), str(width))
                tc_w.set(qn("w:type"), "dxa")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def protect_table(table, repeat_header: bool = True):
    for row in table.rows:
        set_row_cant_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
    if repeat_header and table.rows:
        set_row_repeat_header(table.rows[0])


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, FONT_NUMERIC, 8, False, COLORS["gray_600"])


def clear_story_part(part):
    for child in list(part._element):
        part._element.remove(child)
    part.add_paragraph()


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update field to refresh the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal.font.size = Pt(12)
    normal.font.color.rgb = rgb(COLORS["gray_800"])
    normal.paragraph_format.space_after = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    for name, size, color, before, after in [
        ("Title", 28, COLORS["navy"], 0, 8),
        ("Subtitle", 14, COLORS["gray_600"], 0, 12),
        ("Heading 1", 15, COLORS["navy"], 12, 6),
        ("Heading 2", 14, COLORS["brand"], 10, 6),
        ("Heading 3", 13, COLORS["gray_800"], 8, 6),
    ]:
        style = styles[name]
        style.font.name = FONT_HEADING
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_HEADING)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_HEADING)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["Caption"].font.name = FONT_BODY
    styles["Caption"].font.size = Pt(8.5)
    styles["Caption"].font.color.rgb = rgb(COLORS["gray_600"])
    styles["Caption"].paragraph_format.space_after = Pt(6)
    styles["Caption"].paragraph_format.line_spacing = 1.15

    styles["List Paragraph"].paragraph_format.space_after = Pt(12)
    styles["List Paragraph"].paragraph_format.line_spacing = 1.5

    for style_name in ["TOC Heading", "Intense Quote"]:
        if style_name in styles:
            styles[style_name].font.name = FONT_HEADING


def add_logo(paragraph, width=1.2):
    if LOGO.exists():
        paragraph.add_run().add_picture(str(LOGO), width=Inches(width))
    else:
        run = paragraph.add_run("CQI")
        set_run_font(run, FONT_HEADING, 16, True, COLORS["brand"])


def add_footer(section, doc_title="[Document name]", edit_date="[DD/MM/YY]"):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    clear_story_part(footer)

    line = footer.add_table(rows=1, cols=1, width=Inches(6.77))
    set_table_width(line, CONTENT_DXA, [CONTENT_DXA])
    shade_cell(line.cell(0, 0), COLORS["brand"])
    line.cell(0, 0).height = Pt(2)

    table = footer.add_table(rows=1, cols=2, width=Inches(6.77))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, CONTENT_DXA, [7800, 1944])
    for cell in table.row_cells(0):
        set_cell_margins(cell, 20, 0, 20, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left = table.cell(0, 0).paragraphs[0]
    run = left.add_run(f"{doc_title} · {edit_date}")
    set_run_font(run, FONT_BODY, 7.5, False, COLORS["gray_600"])
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(right)


def add_header(section):
    section.header.is_linked_to_previous = False
    header = section.header
    clear_story_part(header)
    table = header.add_table(rows=1, cols=2, width=Inches(6.77))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, CONTENT_DXA, [7400, 2344])
    for cell in table.row_cells(0):
        set_cell_margins(cell, 0, 0, 0, 0)
    left = table.cell(0, 0).paragraphs[0]
    run = left.add_run("[Document control name]")
    set_run_font(run, FONT_BODY, 7.5, False, COLORS["gray_600"])
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_logo(right, 0.8)


def add_cover(doc: Document):
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_story_part(section.header)
    clear_story_part(section.footer)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.4)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.7)

    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_logo(logo_p, 1.35)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, 44, 0)

    hero = doc.add_table(rows=1, cols=2)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(hero, CONTENT_DXA, [240, 9504])
    shade_cell(hero.cell(0, 0), COLORS["brand"])
    shade_cell(hero.cell(0, 1), COLORS["navy"])
    for cell in hero.row_cells(0):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 340, 420, 340, 280)
    p = hero.cell(0, 1).paragraphs[0]
    run = p.add_run("[DOCUMENT TYPE]")
    set_run_font(run, FONT_NUMERIC, 9, True, COLORS["accent"])
    p.add_run("\n")
    p.add_run("\n")
    run = p.add_run("[Primary document title]")
    set_run_font(run, FONT_HEADING, 24, True, COLORS["white"])
    p.add_run("\n")
    run = p.add_run("[Subtitle or scope statement]")
    set_run_font(run, FONT_BODY, 11.5, False, COLORS["gray_100"])

    doc.add_paragraph()
    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(meta, 6600, [2200, 4400])
    rows = [
        ("Client / audience", "[Client, steering group, or internal audience]"),
        ("Document ID", "[CQI-CLIENT-WORKSTREAM-DOC-001]"),
        ("Version", "[R01 / Draft / Final]"),
        ("Prepared by", "CQI"),
        ("Date", "[YYYY-MM-DD]"),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = meta.rows[idx].cells
        for cell in cells:
            cell_border(cell)
            set_cell_margins(cell, 110, 120, 110, 120)
        shade_cell(cells[0], COLORS["gray_50"])
        r = cells[0].paragraphs[0].add_run(label)
        set_run_font(r, FONT_HEADING, 8.5, True, COLORS["gray_600"])
        r = cells[1].paragraphs[0].add_run(value)
        set_run_font(r, FONT_BODY, 9, False, COLORS["gray_800"])

    bottom = doc.add_paragraph()
    bottom.paragraph_format.space_before = Pt(150)
    r = bottom.add_run("CQI\nExplainable customer-journey intelligence")
    set_run_font(r, FONT_HEADING, 9, True, COLORS["gray_600"])


def add_control_page(doc: Document):
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.6)
    add_header(section)
    add_footer(section)

    doc.add_heading("Document control", level=1)
    p = doc.add_paragraph("Use this page for traceability before stakeholder circulation.")
    set_paragraph_spacing(p, after=12)

    meta = doc.add_table(rows=6, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(meta, CONTENT_DXA, [1950, 2922, 1950, 2922])
    rows = [
        ("Document ID", "[ID]", "Status", "[Draft / Review / Final]"),
        ("Owner", "[Name / role]", "Reviewer", "[Name / role]"),
        ("Approver", "[Name / role]", "Confidentiality", "[Internal / Client confidential]"),
        ("Source scope", "[Data, repo, client, period]", "Quality gate", "[Passed / warning / failed]"),
        ("Supersedes", "[Previous version]", "Next review", "[YYYY-MM-DD]"),
        ("Language", "[Spanish / English / bilingual]", "Prepared by", "CQI"),
    ]
    for ridx, row in enumerate(rows):
        for cidx, text in enumerate(row):
            cell = meta.cell(ridx, cidx)
            cell_border(cell)
            set_cell_margins(cell)
            if cidx in (0, 2):
                shade_cell(cell, COLORS["gray_50"])
                run = cell.paragraphs[0].add_run(text)
                set_run_font(run, FONT_HEADING, 8.5, True, COLORS["gray_600"])
            else:
                run = cell.paragraphs[0].add_run(text)
                set_run_font(run, FONT_BODY, 9, False, COLORS["gray_800"])

    doc.add_heading("Revision history", level=2)
    table = doc.add_table(rows=2, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, CONTENT_DXA, [680, 1120, 3424, 1130, 1130, 1130, 1130])
    headers = ["Rev", "Date", "Description", "Issued by", "Checked by", "Approved by", "Client approval"]
    values = ["R01", "[YYYY-MM-DD]", "[Issued for review]", "[Name]", "[Name]", "[Name]", "[Y/N]"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, COLORS["navy"])
        cell_border(cell, COLORS["navy"])
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, FONT_HEADING, 8, True, COLORS["white"])
    for idx, text in enumerate(values):
        cell = table.cell(1, idx)
        cell_border(cell)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, FONT_BODY, 8.5, False, COLORS["gray_800"])


def add_toc_page(doc: Document):
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    add_header(section)
    add_footer(section)

    doc.add_heading("Index", level=1)
    p = doc.add_paragraph()
    add_toc(p)
    set_paragraph_spacing(p, after=12)
    p = doc.add_paragraph("Note: In Word, right-click the table of contents and choose Update Field before issuing a final version.")
    p.style = "Caption"


def add_callout(doc: Document, label: str, text: str, color: str = "brand"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, CONTENT_DXA, [180, 9564])
    shade_cell(table.cell(0, 0), COLORS[color])
    shade_cell(table.cell(0, 1), COLORS["gray_50"])
    for cell in table.row_cells(0):
        set_cell_margins(cell, 150, 150, 150, 150)
        cell_border(cell, COLORS["gray_200"])
    p = table.cell(0, 1).paragraphs[0]
    r = p.add_run(f"{label}: ")
    set_run_font(r, FONT_HEADING, 9, True, COLORS[color])
    r = p.add_run(text)
    set_run_font(r, FONT_BODY, 9.5, False, COLORS["gray_800"])


def add_apa_title(doc: Document, label: str, title: str):
    number = doc.add_paragraph()
    number.paragraph_format.space_before = Pt(6)
    number.paragraph_format.space_after = Pt(0)
    number.paragraph_format.keep_with_next = True
    run = number.add_run(label)
    set_run_font(run, FONT_BODY, 12, True, COLORS["gray_800"])

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    title_p.paragraph_format.keep_with_next = True
    run = title_p.add_run(title)
    set_run_font(run, FONT_BODY, 12, False, COLORS["gray_800"])
    run.italic = True


def add_apa_note(doc: Document, text: str):
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(12)
    note.paragraph_format.line_spacing = 1.15
    run = note.add_run("Note. ")
    set_run_font(run, FONT_BODY, 9.5, True, COLORS["gray_600"])
    run = note.add_run(text)
    set_run_font(run, FONT_BODY, 9.5, False, COLORS["gray_600"])


def add_body_sample(doc: Document):
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    add_header(section)
    add_footer(section)

    doc.add_heading("Executive summary", level=1)
    p = doc.add_paragraph("[Front-load the conclusion in 4-6 lines. State what changed, why it matters, and what decision is required.]")
    set_paragraph_spacing(p, after=10)
    add_callout(doc, "Decision", "[Recommended action, owner, and decision date.]", "brand")

    doc.add_heading("Key indicators", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, CONTENT_DXA, [2436, 2436, 2436, 2436])
    labels = ["Repair pressure", "Risk pressure", "Operational score", "Evidence status"]
    values = ["[00.0]", "[00.0]", "[00/100]", "[Validation-pending]"]
    colors = ["repair", "risk", "operational", "brand"]
    for idx, label in enumerate(labels):
        cell = table.cell(0, idx)
        shade_cell(cell, COLORS["gray_50"])
        cell_border(cell)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(label)
        set_run_font(run, FONT_HEADING, 8, True, COLORS["gray_600"])
        cell = table.cell(1, idx)
        cell_border(cell)
        set_cell_margins(cell, 160, 120, 160, 120)
        run = cell.paragraphs[0].add_run(values[idx])
        set_run_font(run, FONT_NUMERIC, 14, True, COLORS[colors[idx]])

    doc.add_heading("1. Context and scope", level=1)
    doc.add_paragraph("[Describe the client, workstream, source systems, period covered, and business question. Include exact data grain and version boundaries when relevant.]")
    doc.add_heading("1.1 Source lineage", level=2)
    add_callout(doc, "Lineage", "[Source path/table/API] -> [processing step] -> [deliverable section]. Separate verified facts from assumptions.", "repair")

    doc.add_heading("2. Findings", level=1)
    doc.add_paragraph("[Use concise prose. Every figure/table must be numbered, referenced in the text, and explained.]")
    doc.add_heading("2.1 Evidence table pattern", level=2)
    add_apa_title(doc, "Table 1", "Evidence Matrix With CQI Table Geometry")
    table = doc.add_table(rows=4, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, CONTENT_DXA, [1100, 2161, 2161, 2161, 2161])
    headers = ["ID", "Finding", "Evidence", "Impact", "Action"]
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, COLORS["navy"])
        cell_border(cell, COLORS["navy"])
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, FONT_HEADING, 8, True, COLORS["white"])
    for ridx in range(1, 4):
        values = [f"F{ridx}", "[Finding]", "[Source / metric]", "[Business impact]", "[Owner / next step]"]
        for cidx, text in enumerate(values):
            cell = table.cell(ridx, cidx)
            cell_border(cell)
            set_cell_margins(cell)
            if cidx == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cell.paragraphs[0].add_run(text)
                set_run_font(run, FONT_NUMERIC, 8.5, True, COLORS["brand"])
            else:
                run = cell.paragraphs[0].add_run(text)
                set_run_font(run, FONT_BODY, 8.5, False, COLORS["gray_800"])
    add_apa_note(doc, "Replace this note with source, scope, sample size, caveat, or calculation boundary.")

    doc.add_heading("2.2 Figure pattern", level=2)
    add_apa_title(doc, "Figure 1", "Journey Pressure Example Placeholder")
    figure = doc.add_table(rows=1, cols=1)
    figure.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(figure, CONTENT_DXA, [CONTENT_DXA])
    cell = figure.cell(0, 0)
    shade_cell(cell, COLORS["gray_50"])
    cell_border(cell, COLORS["gray_200"])
    set_cell_margins(cell, 520, 520, 520, 520)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Insert figure, chart, or image here. Keep within document margins.]")
    set_run_font(run, FONT_BODY, 12, False, COLORS["gray_600"])
    add_apa_note(doc, "Every figure must be referenced in the text and explained before or immediately after it.")

    doc.add_heading("3. Recommendations", level=1)
    doc.add_paragraph("[Prioritize actions. For each recommendation, include owner, dependency, confidence, and validation caveat.]")
    add_callout(doc, "Guardrail", "Operational scores and churn labels must state whether they are scoring inputs or validation-only evidence.", "risk")

    doc.add_page_break()
    doc.add_heading("Appendix A. Method and definitions", level=1)
    doc.add_paragraph("[Definitions, formulas, assumptions, exclusions, source extracts, and validation notes.]")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_control_page(doc)
    add_toc_page(doc)
    add_body_sample(doc)

    props = doc.core_properties
    props.title = "CQI Document Template"
    props.subject = "Canonical Word template aligned to the CQI Design System"
    props.author = "CQI"
    props.keywords = "CQI, Design System, Word template"
    for table in doc.tables:
        protect_table(table, repeat_header=len(table.rows) >= 2)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
